'''
create an interactive cytometry plot

initialise
configure
update data
'''

import sys
from PySide6 import QtWidgets, QtCore, QtGui
import pyqtgraph as pg
import flowkit as fk
import time
import numpy as np
import colorcet as cc
from PySide6.QtCore import QPoint
from PySide6.QtGui import QPixmap, Qt, QPainter

transforms_menu_items = ['Linear', 'Logicle', 'Log']
label_offset = (0, -0.03)

def export_docx(image_file):
    import docx
    from docx.shared import Inches

    # Simple usage - just 3 lines to create a document with an image!
    doc = docx.Document()
    doc.add_heading('My Document', 0)
    doc.add_paragraph('Hello World!')

    # Add an image - one line!
    doc.add_picture(image_file, width=Inches(6))

    # Save - one line!
    doc.save('document.docx')

def export_high_quality(widget, filename, scale_factor=2.0):
    """Export widget with higher resolution"""
    original_size = widget.size()

    # Create larger pixmap for better quality
    scaled_size = original_size * scale_factor
    pixmap = QPixmap(scaled_size)
    pixmap.fill(Qt.white)

    painter = QPainter(pixmap)
    painter.scale(scale_factor, scale_factor)
    widget.render(painter, QPoint(0, 0))
    painter.end()

    # Save scaled down if needed
    pixmap.save(filename)
    print(f"High-quality export to {filename}")

def render_widget_light_pg(widget, filename, scale_factor=2):
    # --- Save old PyQtGraph colors ---
    old_bg = pg.getConfigOption("background")
    old_fg = pg.getConfigOption("foreground")
    print([old_bg, old_fg])

    # --- Set temporary light theme ---
    pg.setConfigOption("background", "w")   # white
    pg.setConfigOption("foreground", "k")   # black

    # Apply to all contained plots
    widget.repaint()

    # --- Render the widget ---
    pm = QPixmap(widget.size()*scale_factor)
    pm.fill(Qt.transparent)
    painter = QPainter(pm)
    painter.scale(scale_factor, scale_factor)
    widget.render(painter, QPoint(0, 0))
    painter.end()

    # --- Restore original PyQtGraph colors ---
    pg.setConfigOption("background", old_bg)
    pg.setConfigOption("foreground", old_fg)
    widget.repaint()

    pm.save(filename)
    print(f"High-quality export to {filename}")

    return pm


def export_widget_as_image(widget, filename):
    """Export any widget as an image"""
    pixmap = QPixmap(widget.size())
    widget.render(pixmap)
    pixmap.save(filename)
    print(f"Widget exported as {filename}")


class LabelEditDialog(QtWidgets.QDialog):
    """Minimal dialog with just a QLineEdit for editing text."""
    def __init__(self, text="", existing_names=[], parent=None):
        super().__init__(parent)
        self.setWindowTitle("Edit Label")
        self.setModal(True)
        self.setFixedSize(250, 80)
        self.existing_names = existing_names
        self.old_name = text

        layout = QtWidgets.QVBoxLayout(self)

        self.line_edit = QtWidgets.QLineEdit(self)
        self.line_edit.setText(text)
        self.line_edit.selectAll()   # auto-select text
        self.line_edit.setFocus()    # auto-focus
        layout.addWidget(self.line_edit)

        buttons = QtWidgets.QDialogButtonBox(
            QtWidgets.QDialogButtonBox.StandardButton.Ok | QtWidgets.QDialogButtonBox.StandardButton.Cancel,
            QtCore.Qt.Orientation.Horizontal, self
        )
        buttons.accepted.connect(self.validate_and_accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)

    def validate_and_accept(self):
        text = self.line_edit.text().strip()
        if not text:
            QtWidgets.QMessageBox.warning(self, "Error", "Input cannot be empty.")
        elif text == self.old_name:
            self.reject()
        elif text in self.existing_names:
            QtWidgets.QMessageBox.warning(self, "Error", f'"{text}" already exists.')
        else:
            self.accept()

    def getText(self):
        return self.line_edit.text()


class DraggableRoiLabel(pg.TextItem):
    labelRenamed = QtCore.Signal(str)
    """A text label that can be dragged and renamed via small dialog."""
    def __init__(self, parent_roi, text, gates, pos=(0,0), anchor=(0, 0)):
        super().__init__(text, anchor=anchor)
        self.parent_roi = parent_roi
        self.gates = gates
        self.setPos(*pos)

        # Style: bold + bigger font
        font = QtGui.QFont("Arial", 12, QtGui.QFont.Weight.Bold)
        self.setFont(font)
        self.setColor("k")
        self.fill = pg.mkBrush(0, 255, 0, 128)

        self.setFlag(self.GraphicsItemFlag.ItemIsMovable, True)
        self.setFlag(self.GraphicsItemFlag.ItemIsSelectable, True)

    def paint(self, p, *args):
        """Draw background behind text."""
        rect = self.boundingRect().adjusted(-2, -2, 2, 2)
        p.setBrush(self.fill)
        p.setPen(pg.mkPen("k"))
        p.drawRect(rect)
        super().paint(p, *args)

    def mouseReleaseEvent(self, event):
        super().mouseReleaseEvent(event)
        roi_pos = self.parent_roi.pos()
        self.parent_roi.label_offset = (
            self.pos().x() - roi_pos.x(),
            self.pos().y() - roi_pos.y()
        )

    def mouseDoubleClickEvent(self, event):
        """Open small dialog for editing text."""
        gate_names = ['root'] + [g[0] for g in self.gates.get_gate_ids()]
        dlg = LabelEditDialog(self.toPlainText(), existing_names=gate_names)
        if dlg.exec() == QtWidgets.QDialog.DialogCode.Accepted:
            new_text = dlg.getText()
            if new_text.strip():
                self.setText(new_text)
                self.labelRenamed.emit(new_text)


class Transform:
    def __init__(self, linear_t=262144, linear_a=100, logicle_t=262144, logicle_w=0.5, logicle_m=4.5, logicle_a=0, log_t=262144, log_m=6, scale_bins=200):
        self.linear_xform = None
        self.logicle_xform = None
        self.xform = None
        self.scale = None
        self.step_scale = None
        self.zero_inverse = None
        self.zero = None
        self.ticks = None
        self.linear_t = linear_t
        self.linear_a = linear_a
        self.logicle_t = logicle_t
        self.logicle_w = logicle_w
        self.logicle_m = logicle_m
        self.logicle_a = logicle_a
        self.log_t = log_t
        self.log_m = log_m
        self.scale_bins = scale_bins
        self.initialise_xforms()

    def initialise_xforms(self):
        self.linear_xform = fk.transforms.LinearTransform(param_t=self.linear_t, param_a=self.linear_a)
        self.logicle_xform = fk.transforms.LogicleTransform(param_t=self.logicle_t, param_w=self.logicle_w, param_m=self.logicle_m, param_a=self.logicle_a)
        self.log_xform = fk.transforms.LogTransform(param_t=self.log_t, param_m=self.log_m)

    def set_transform(self, id, limits):
        if id == 0: #'linear'
            self.set_linear(limits)
        elif id == 1: #'logicle'
            self.set_logicle(limits)
        elif id == 2: #'log'
            self.set_log(limits)
        elif id == 'default':
            self.set_default(limits)

    def set_linear(self, limits):
        self.xform = self.linear_xform
        self.scale = np.concatenate((
            [-np.inf],
            self.xform.inverse(np.linspace(limits[0], limits[1], self.scale_bins)),
            [np.inf]
        ))
        self.step_scale = np.concatenate((
            [limits[0]-1/self.scale_bins],
            np.linspace(limits[0], limits[1], self.scale_bins),
            [limits[1]+1/self.scale_bins, limits[1]+2/self.scale_bins]
        ))
        self.zero_inverse = self.xform.inverse(np.array([0]))[0]
        self.zero = self.xform.apply(np.array([0]))[0]
        self.ticks = self.linear_ticks

    def set_logicle(self, limits):
        self.xform = self.logicle_xform
        self.scale = np.concatenate((
            [-np.inf],
            self.xform.inverse(np.linspace(limits[0], limits[1], self.scale_bins)),
            [np.inf]
        ))
        self.step_scale = np.concatenate((
            [limits[0]-1/self.scale_bins],
            np.linspace(limits[0], limits[1], self.scale_bins),
            [limits[1]+1/self.scale_bins, limits[1]+2/self.scale_bins]
        ))
        self.zero_inverse = self.xform.inverse(np.array([0]))[0]
        self.zero = self.xform.apply(np.array([0]))[0]
        self.ticks = self.logicle_ticks

    def set_log(self, limits):
        self.xform = self.log_xform
        self.scale = np.concatenate((
            [-np.inf],
            self.xform.inverse(np.linspace(limits[0], limits[1], self.scale_bins)),
            [np.inf]
        ))
        self.step_scale = np.concatenate((
            [limits[0]-1/self.scale_bins],
            np.linspace(limits[0], limits[1], self.scale_bins),
            [limits[1]+1/self.scale_bins, limits[1]+2/self.scale_bins]
        ))
        self.zero_inverse = -np.inf
        self.zero = self.xform.apply(np.array([1]))[0]
        self.ticks = self.log_ticks

    def set_default(self, limits):
        self.xform = None
        self.scale = np.concatenate((
            [-np.inf],
            np.arange(0,self.scale_bins),
            [np.inf]
        ))
        self.zero_inverse = 0
        self.zero = 0
        self.ticks = self.default_ticks

    def logicle_ticks(self,axis):
        major_values = np.concatenate([
            -np.logspace(3, 0, 4),  # Negative values
            [0],
            np.logspace(0, 6, 7)  # Positive values
        ])
        minor_values = np.hstack([m * np.arange(0.1, 1, 0.1) if m != 0 else None for m in major_values])

        # Transform to plot coordinates
        trans_major_values = self.logicle_xform.apply(major_values)
        trans_minor_values = self.logicle_xform.apply(minor_values)

        superscripts = {
            "0": "⁰", "1": "¹", "2": "²", "3": "³", "4": "⁴",
            "5": "⁵", "6": "⁶", "7": "⁷", "8": "⁸", "9": "⁹", "-": "⁻"
        }

        def to_superscript(n):
            return "".join(superscripts.get(c, c) for c in str(n))

        # Create ticks with formatted labels
        major_ticks = []
        for val, raw in zip(trans_major_values, major_values):
            if raw == 0:
                label = "0"
            elif abs(raw) < 100:
                label = ""
            elif raw < 0:
                label = f"-10{to_superscript(int(np.log10(-raw)))}" if abs(raw) >= 10 else f"{raw:.1f}"
            else:
                label = f"10{to_superscript(int(np.log10(raw)))}" if raw >= 10 else f"{raw:.1f}"
            major_ticks.append((val, label))

        minor_ticks = [(val, '') for val in trans_minor_values]

        axis.setTicks([minor_ticks, major_ticks])

    def log_ticks(self,axis):
        major_values = np.logspace(0, 6, 7)
        minor_values = np.hstack([m * np.arange(0.1, 1, 0.1) if m != 0 else None for m in major_values])

        # Transform to plot coordinates
        trans_major_values = self.log_xform.apply(major_values)
        trans_minor_values = self.log_xform.apply(minor_values)

        superscripts = {
            "0": "⁰", "1": "¹", "2": "²", "3": "³", "4": "⁴",
            "5": "⁵", "6": "⁶", "7": "⁷", "8": "⁸", "9": "⁹", "-": "⁻"
        }

        def to_superscript(n):
            return "".join(superscripts.get(c, c) for c in str(n))

        # Create ticks with formatted labels
        major_ticks = []
        for val, raw in zip(trans_major_values, major_values):
            label = f"10{to_superscript(int(np.log10(raw)))}" if raw >= 10 else f"{raw:.1f}"
            major_ticks.append((val, label))
        minor_ticks = [(val, '') for val in trans_minor_values]
        axis.setTicks([minor_ticks, major_ticks])

    def linear_ticks(self,axis):
        axis.setTicks(None)

        major_values = np.arange(0,1e6,1e5)
        minor_values = np.arange(0,1e6,1e4)

        # Transform to plot coordinates
        trans_major_values = self.linear_xform.apply(major_values)
        trans_minor_values = self.linear_xform.apply(minor_values)

        superscripts = {
            "0": "⁰", "1": "¹", "2": "²", "3": "³", "4": "⁴",
            "5": "⁵", "6": "⁶", "7": "⁷", "8": "⁸", "9": "⁹", "-": "⁻"
        }

        def to_superscript(n):
            return "".join(superscripts.get(c, c) for c in str(n))

        # Create ticks with formatted labels
        major_ticks = []
        for val, raw in zip(trans_major_values, major_values):
            if raw == 0:
                label = "0"
            elif abs(raw) < 100:
                label = ""
            elif raw < 0:
                label = f"-10{to_superscript(int(np.log10(-raw)))}" if abs(raw) >= 10 else f"{raw:.1f}"
            else:
                r,i = np.modf(np.log10(raw))
                i = int(i)
                r = 10**r
                label = f"{r:.0f}.10{to_superscript(i)}" if raw >= 10 else f"{raw:.1f}"
            major_ticks.append((val, label))

        minor_ticks = [(val, '') for val in trans_minor_values]

        axis.setTicks([minor_ticks, major_ticks])

    def default_ticks(self,axis):
        axis.setTicks(None)


class ContextMenuTargetItem(pg.TargetItem):
    # Optional: custom signal for removal
    sigRemoveRequested = QtCore.Signal(object)

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)

        # Create context menu
        self.menu = QtWidgets.QMenu()
        self.action_remove = self.menu.addAction("Remove Region")

    def mouseClickEvent(self, ev):
        if ev.button() == QtCore.Qt.MouseButton.RightButton:
            self.menu.exec(ev.screenPos().toPoint())
            ev.accept()
        else:
            super().mouseClickEvent(ev)


class QuadROI(pg.ROI):
    sigPosChanged = QtCore.Signal(float, float)  # min_x, max_x

    def __init__(self, vb, x=10, y=50, pen=pg.mkPen('g', width=3), handle_size=15, *args, **kwargs):
        super().__init__((0,0), *args, **kwargs)
        self.vx = pg.InfiniteLine(pos=x, angle=90, movable=True, pen=pen)
        self.vy = pg.InfiniteLine(pos=y, angle=0, movable=True, pen=pen)

        # Connect movement signals
        self.vx.sigPositionChanged.connect(self._line_moved)
        self.vy.sigPositionChanged.connect(self._line_moved)
        self.vb = vb
        self.vb.addItem(self.vx)
        self.vb.addItem(self.vy)
        self.addFreeHandle([0, 0])
        self.label = None
        self.gate_name = None

        self.target = ContextMenuTargetItem(pos = (x, y), size = 20, pen = 'c', symbol = 's')
        self.vb.addItem(self.target)
        self.target.sigPositionChanged.connect(self._target_moved)
        self.target.action_remove.triggered.connect(self.request_remove)

    def _line_moved(self, line):
        # Update region when a line moves
        x = self.vx.value()
        y = self.vy.value()
        self.target.setPos((x,y))
        self.sigPosChanged.emit(x, y)

    def _target_moved(self, target):
        # Update region when target moves
        pos = target.pos()
        x = pos.x()
        y = pos.y()
        self.vx.setPos(x)
        self.vy.setPos(y)
        self.sigPosChanged.emit(x, y)

    def pos(self):
        x = self.vx.value()
        y = self.vy.value()
        return QtCore.QPointF(x,y)

    def request_remove(self):
        """Emit sigRemoveRequested (same as built-in ROI behavior)."""
        self.vb.removeItem(self.vx)
        self.vb.removeItem(self.vy)
        self.vb.removeItem(self.target)
        self.sigRemoveRequested.emit(self)


class ContextMenuRangeROI(pg.LinearRegionItem):
    # Optional: custom signal for removal
    sigRemoveRequested = QtCore.Signal(object)

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)

        # Create context menu
        self.menu = QtWidgets.QMenu()
        self.action_remove = self.menu.addAction("Remove Region")

    def mouseClickEvent(self, ev):
        if ev.button() == QtCore.Qt.MouseButton.RightButton:
            self.menu.exec(ev.screenPos().toPoint())
            ev.accept()
        else:
            super().mouseClickEvent(ev)


class RangeROI(pg.ROI):
    sigRangeChanged = QtCore.Signal(float, float)  # min_x, max_x

    def __init__(self, vb, x1=10, x2=50, pen=pg.mkPen('g', width=3), handle_size=15, *args, **kwargs):
        super().__init__((0,0), *args, **kwargs)
        self.v1 = pg.InfiniteLine(pos=x1, angle=90, movable=True, pen=pen)
        self.v2 = pg.InfiniteLine(pos=x2, angle=90, movable=True, pen=pen)

        # Connect movement signals
        self.v1.sigPositionChanged.connect(self._line_moved)
        self.v2.sigPositionChanged.connect(self._line_moved)
        self.region = ContextMenuRangeROI(values=(x1, x2))
        self.vb = vb
        self.vb.addItem(self.region)
        self.vb.addItem(self.v1)
        self.vb.addItem(self.v2)
        self.region.sigRegionChanged.connect(self._region_moved)

        self.label = None
        self.gate_name = None
        self.label_offset = label_offset

        # Connect actions
        self.region.action_remove.triggered.connect(self.request_remove)

    def _line_moved(self, line):
        # Update region when a line moves
        min_x = min(self.v1.value(), self.v2.value())
        max_x = max(self.v1.value(), self.v2.value())
        self.region.setRegion((min_x, max_x))
        self.sigRangeChanged.emit(min_x, max_x)

    def _region_moved(self):
        # Update lines if the region is dragged
        min_x, max_x = self.region.getRegion()
        self.v1.setPos(min_x)
        self.v2.setPos(max_x)
        self.sigRangeChanged.emit(min_x, max_x)

    def setRange(self, x1, x2):
        self.v1.setPos(x1)
        self.v2.setPos(x2)
        self.region.setRegion((x1, x2))
        self.sigRangeChanged.emit(x1, x2)

    def pos(self):
        min_x, max_x = self.region.getRegion()
        return QtCore.QPointF(min_x,1)

    def request_remove(self):
        """Emit sigRemoveRequested (same as built-in ROI behavior)."""
        self.vb.removeItem(self.region)
        self.vb.removeItem(self.v1)
        self.vb.removeItem(self.v2)
        self.sigRemoveRequested.emit(self)


class BigHandlesPolyLineROI(pg.PolyLineROI):
    def __init__(self, *args, **kwargs):
        self.handleSize = 15  # Set before super().__init__()
        super().__init__(*args, **kwargs)
        self.label = None
        self.gate_name = None
        self.label_offset = label_offset

        # Create context menu
        self.menu = QtWidgets.QMenu()
        self.action_delete = self.menu.addAction("Delete Gate", self.request_remove)

    def addHandle(self, *args, **kwargs):
        self.handleSize = 15
        return super().addHandle(*args, **kwargs)

    def mouseClickEvent(self, event):
        if event.button() == QtCore.Qt.MouseButton.RightButton:
            # Show context menu at cursor position
            self.menu.exec(event.screenPos().toPoint())
            event.accept()
        else:
            # Keep normal ROI drag/resize behavior
            super().mouseClickEvent(event)

    def request_remove(self):
        """Emit sigRemoveRequested (same as built-in ROI behavior)."""
        self.sigRemoveRequested.emit(self)


class BigHandlesRectROI(pg.RectROI):
    def __init__(self, *args, **kwargs):
        self.handleSize = 15  # Set before super().__init__()
        super().__init__(*args, **kwargs)
        self.label = None
        self.gate_name = None
        self.label_offset = label_offset

        # Create context menu
        self.menu = QtWidgets.QMenu()
        self.action_delete = self.menu.addAction("Delete Gate", self.request_remove)

    def addHandle(self, *args, **kwargs):
        self.handleSize = 15
        return super().addHandle(*args, **kwargs)

    def mouseClickEvent(self, event):
        if event.button() == QtCore.Qt.MouseButton.RightButton:
            # Show context menu at cursor position
            self.menu.exec(event.screenPos().toPoint())
            event.accept()
        else:
            # Keep normal ROI drag/resize behavior
            super().mouseClickEvent(event)

    def request_remove(self):
        """Emit sigRemoveRequested (same as built-in ROI behavior)."""
        self.sigRemoveRequested.emit(self)


class BigHandlesEllipseROI(pg.EllipseROI):
    def __init__(self, *args, **kwargs):
        self.handleSize = 15  # Set before super().__init__()
        super().__init__(*args, **kwargs)
        self.label = None
        self.gate_name = None
        self.label_offset = label_offset

        # Create context menu
        self.menu = QtWidgets.QMenu()
        self.action_delete = self.menu.addAction("Delete Gate", self.request_remove)

    def addHandle(self, *args, **kwargs):
        self.handleSize = 15
        return super().addHandle(*args, **kwargs)

    def mouseClickEvent(self, event):
        if event.button() == QtCore.Qt.MouseButton.RightButton:
            # Show context menu at cursor position
            self.menu.exec(event.screenPos().toPoint())
            event.accept()
        else:
            # Keep normal ROI drag/resize behavior
            super().mouseClickEvent(event)

    def request_remove(self):
        """Emit sigRemoveRequested (same as built-in ROI behavior)."""
        self.sigRemoveRequested.emit(self)


class ZoomAxis(pg.AxisItem):
    def __init__(self, interactive_cytometry_plot, orientation, viewbox, **kwargs):
        super().__init__(orientation, **kwargs)
        self.setAcceptHoverEvents(True)  # needed for hover detection
        self.vb = viewbox
        self.orientation = orientation
        self.initial_pos = None
        self._last_pos = None
        self._pending_delta = 0
        self.zoom_timer = QtCore.QTimer()
        self.zoom_timer.setInterval(60)  # update rate in ms
        self.zoomZero = 0
        self.fullRange = (0, 1.1)
        self.limits = (0, 1)
        self.interactive_cytometry_plot = interactive_cytometry_plot

    def hoverEnterEvent(self, event):
        QtWidgets.QApplication.setOverrideCursor(QtCore.Qt.CursorShape.PointingHandCursor)

    def hoverLeaveEvent(self, event):
        QtWidgets.QApplication.restoreOverrideCursor()

    def mousePressEvent(self, ev):
        if ev.button() == QtCore.Qt.MouseButton.LeftButton:
            self.initial_pos = ev.pos()
            self._last_pos = self.initial_pos
            self._pending_delta = 0
            self.zoom_timer.start()
            ev.accept()
        else:
            ev.ignore()

    def mouseMoveEvent(self, ev):
        if self._last_pos is None:
            ev.ignore()
            return

        delta = ev.pos() - self._last_pos
        self._last_pos = ev.pos()

        if self.orientation == 'bottom':
            self._pending_delta += delta.x()
        elif self.orientation == 'left':
            self._pending_delta += delta.y()

        ev.accept()

        if QtWidgets.QApplication.overrideCursor() == QtCore.Qt.CursorShape.PointingHandCursor:
            QtWidgets.QApplication.setOverrideCursor(QtCore.Qt.CursorShape.ClosedHandCursor)


    def mouseReleaseEvent(self, ev):
        self._last_pos = None
        self._pending_delta = 0
        self.zoom_timer.stop()
        ev.accept()

        QtWidgets.QApplication.restoreOverrideCursor()


class NoPanViewBox(pg.ViewBox):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, enableMouse=True, **kwargs)
        self.setMouseMode(self.PanMode)

    def mouseDragEvent(self, ev, axis=None):
        ev.ignore()  # 🔒 Disable all panning and dragging in plot area

    def wheelEvent(self, ev, axis=None):
        ev.ignore()  # 🔒 Disable mouse wheel zoom


class InteractiveLabel(pg.LabelItem):
    def __init__(self, text="", angle=0, **kwargs):
        super().__init__(text, angle=angle, **kwargs)
        self.setAcceptHoverEvents(True)
        self._default_font = self.item.font()
        self._hover_font = QtGui.QFont(self._default_font)
        self._hover_font.setUnderline(True)
        self.leftClickMenuItems = []
        self.rightClickMenuItems = []
        self.leftItemSelected = None
        self.rightItemSelected = None
        self.leftClickMenuFunction = None
        self.rightClickMenuFunction = None

    def hoverEnterEvent(self, event):
        self.item.setFont(self._hover_font)
        QtWidgets.QApplication.setOverrideCursor(QtCore.Qt.CursorShape.PointingHandCursor)

    def hoverLeaveEvent(self, event):
        self.item.setFont(self._default_font)
        QtWidgets.QApplication.restoreOverrideCursor()

    def mousePressEvent(self, event):
        if event.button() == QtCore.Qt.MouseButton.RightButton:
            self.show_right_context_menu(event.screenPos())
        elif event.button() == QtCore.Qt.MouseButton.LeftButton:
            self.selectable_menu_activates_function(event.screenPos())  # can choose to differentiate

    def selectable_menu_activates_function(self, pos):
        menu = QtWidgets.QMenu()
        actions = {}

        for n, item in enumerate(self.leftClickMenuItems):
            if item == 'Count':
                menu.addSeparator()
                action = menu.addAction(item)
                actions[action] = item
            elif item == 'All Fluorescence':
                menu.addSeparator()
                action = menu.addAction(item)
                actions[action] = item
            else:
                action = menu.addAction(item)
                action.setCheckable(True)
                if n == self.leftItemSelected:
                    action.setChecked(True)
                actions[action] = n
        chosen_action = menu.exec(pos)

        if chosen_action:
            self.leftItemSelected = actions[chosen_action]
            func = self.leftClickMenuFunction
            func(self.leftItemSelected)
            print(f"{func} called with option {self.leftItemSelected}")


    def show_right_context_menu(self, pos):
        menu = QtWidgets.QMenu()
        actions = {}

        for n, item in enumerate(self.rightClickMenuItems):
            action = menu.addAction(item)
            action.setCheckable(True)
            if n == self.rightItemSelected:
                action.setChecked(True)
            actions[action] = n
        chosen_action = menu.exec(pos)

        if chosen_action:
            self.rightItemSelected = actions[chosen_action]
            func = self.rightClickMenuFunction
            func(self.rightItemSelected)
            print(f"{func} called with option {self.rightItemSelected}")



class InteractiveCytometryPlotWidget(QtWidgets.QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)

        # Create main layout
        main_layout = QtWidgets.QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)

        # # Create control buttons layout... placeholder for buttons
        # control_layout = QtWidgets.QHBoxLayout()

        # Access the internal layout (QGraphicsGridLayout)
        self.graphics_widget = pg.GraphicsLayoutWidget()
        layout = self.graphics_widget.ci.layout

        # Set spacing between items (in pixels)
        layout.setHorizontalSpacing(0)
        layout.setVerticalSpacing(0)

        # Add plot title, ViewBox for plotting,
        self.plot_title = InteractiveLabel("Plot Title")
        self.graphics_widget.addItem(self.plot_title, row=0, col=2)
        self.vb = NoPanViewBox()
        self.graphics_widget.addItem(self.vb, row=1, col=2)

        # Y axis label, Y axis itself, X axis, X axis label, Link axes to viewbox
        self.label_y = InteractiveLabel("Y Axis", angle=-90)
        self.graphics_widget.addItem(self.label_y, row=1, col=0)
        self.axis_left = ZoomAxis(self,'left', self.vb)
        self.graphics_widget.addItem(self.axis_left, row=1, col=1)
        self.axis_bottom = ZoomAxis(self,'bottom', self.vb)
        self.graphics_widget.addItem(self.axis_bottom, row=2, col=2)
        self.label_x = InteractiveLabel("X Axis")
        self.graphics_widget.addItem(self.label_x, row=3, col=2)
        self.axis_left.linkToView(self.vb)
        self.axis_bottom.linkToView(self.vb)
        # connect applyZoom method
        self.axis_bottom.zoom_timer.timeout.connect(lambda: self.apply_zoom('x'))
        self.axis_left.zoom_timer.timeout.connect(lambda: self.apply_zoom('y'))

        # connect right click menu
        self.vb.raiseContextMenu = self.right_click_menu

        # initialise configuration
        self.channel_labels = []
        self.scatter_indices = []
        self.fluoro_indices = []
        self.transform_map = []
        self.limits_map = []
        self.gates = []

        # set default channels
        self.id_channel_x = 0
        self.id_channel_y = 'Count'

        # initialise transform objects
        self.transform_x = Transform()
        self.transform_y = Transform()

        # set default source gate
        self.id_source_gate = 0 # i.e. root

        # initialise image and histogram curve
        self.event_data_all_channels = None
        self.img = pg.ImageItem()
        self.hist = pg.PlotDataItem(stepMode='center', fillLevel=0, brush=(100, 100, 250, 150))
        self.count = None

        # initialise dict of rois - key is roi, value is gate id
        self.rois = {}

        # Add widgets to main layout
        # main_layout.addLayout(control_layout) # placeholder for extra buttons
        main_layout.addWidget(self.graphics_widget)


    def configure_plot(self, channel_labels, scatter_indices, fluoro_indices, transform_map, limits_map, gates):

        # store configuration
        self.channel_labels = channel_labels
        self.scatter_indices = scatter_indices
        self.fluoro_indices = fluoro_indices
        self.transform_map = transform_map
        self.limits_map = limits_map
        self.gates = gates

        # initialise transform objects
        self.transform_x.initialise_xforms()
        self.transform_y.initialise_xforms()

        # set default transforms
        if len(self.transform_map) == 0:
            self.transform_map = {}
            for n, label in enumerate(self.channel_labels):
                if n in self.scatter_indices:
                    self.transform_map[n] = 0
                elif n in self.fluoro_indices:
                    self.transform_map[n] = 1
                else:
                    self.transform_map[n] = 'default'
        self.transform_map['Count'] = 'default'
        self.transform_map['All Fluorescence'] = 'default'
        self.transform_map['Intensity'] = 1

        # set default limits
        if len(self.limits_map) == 0:
            self.limits_map = {}
            for n, label in enumerate(self.channel_labels):
                self.limits_map[n] = (0,1)
        self.limits_map['Count'] = (0,100)
        self.limits_map['All Fluorescence'] = (0,len(self.fluoro_indices))
        self.limits_map['Intensity'] = (0,1)

        # initialise image and histogram curve
        self.configure_hist2d()
        self.configure_hist1d()

        # initialise axes, labels and menus
        self.configure_axes_labels_menus()

        # initialise dict of rois - key is roi, value is gate id
        self.rois = {}

    def right_click_menu(self, ev):
        """Override the default context menu."""
        menu = QtWidgets.QMenu()
        if self.id_channel_x == 'All Fluorescence':
            pass
        elif self.id_channel_y == 'Count':
            menu.addAction("New Range Gate", self.new_range_gate)
        else:
            menu.addAction("New Polygon Gate", self.new_polygon_gate)
            menu.addAction("New Rectangle Gate", self.new_rectangle_gate)
            menu.addAction("New Ellipse Gate", self.new_ellipse_gate)
            menu.addAction("New Quadrant Gate", self.new_quadrant_gate)

        menu.addSeparator()
        menu.addAction("Reset Axes Limits", self.vb.autoRange)
        menu.exec(ev.screenPos().toPoint())

    def new_polygon_gate(self):
        # Will hold our polygon ROI
        self.poly_roi = None
        self.vertex_plot = pg.ScatterPlotItem(
            pen='r', brush='g', size=15, symbol='s'
        )
        self.vb.addItem(self.vertex_plot)

        # Connect mouse events
        self.vertices = []
        self.vb.setCursor(QtCore.Qt.CursorShape.CrossCursor)
        self.vb.scene().sigMouseClicked.connect(self.drawing_polygon_handle_click)

    def drawing_polygon_handle_click(self, event):
        pos = self.vb.mapSceneToView(event.scenePos())
        print(pos)
        if event.button() == QtCore.Qt.MouseButton.LeftButton:
            # Add vertex to current polygon
            self.vertices.append(pos)
            self.drawing_polygon_update_vertex_display()

            if event._double:
                if len(self.vertices) >= 4:
                    # Finish polygon
                    # remove last point which is duplicated
                    self.vertices = self.vertices[:-1]
                    # remove another point if it coincides with first point
                    length = QtCore.QLineF(self.vertices[0], self.vertices[-1]).length()
                    if length < 0.1:
                        self.vertices = self.vertices[:-1]

                    self.create_polygon_roi_and_gate()
                    self.vb.setCursor(QtCore.Qt.CursorShape.ArrowCursor)
                else:
                    print("Need at least 3 vertices to create polygon")

    def drawing_polygon_update_vertex_display(self):
        """Show the vertices being added"""
        xs = [v.x() for v in self.vertices]
        ys = [v.y() for v in self.vertices]
        self.vertex_plot.setData(x=xs, y=ys)

        # Draw connecting lines
        if hasattr(self, 'connector'):
            self.vb.removeItem(self.connector)
        self.connector = pg.PlotCurveItem(
            x=xs, y=ys,
            pen=pg.mkPen('y', width=2, style=QtCore.Qt.PenStyle.DashLine)
        )
        self.vb.addItem(self.connector)

    def create_polygon_roi_and_gate(self):
        # Remove temporary connector
        self.vb.scene().sigMouseClicked.disconnect(self.drawing_polygon_handle_click)
        if hasattr(self, 'connector'):
            self.vb.removeItem(self.connector)
            self.vb.removeItem(self.vertex_plot)

        # set name
        name_increment = 1
        while f'P{name_increment}' in [g[0] for g in self.gates.get_gate_ids()]:
            name_increment += 1
        gate_name = f'P{name_increment}'

        # Create PolyLineROI
        roi = BigHandlesPolyLineROI(
            positions=self.vertices,
            closed=True,
            pen=pg.mkPen('g', width=3),
            movable=True,
            removable=True
        )
        roi.gate_name = gate_name
        roi.sigRemoveRequested.connect(self.remove_gate)
        self.vb.addItem(roi)

        # create label
        vertices = np.array([np.array([v.x(), v.y()]) for v in self.vertices])
        xmin = vertices[:,0].min()
        ymin = vertices[:,1].min()
        roi.label_offset = (xmin + label_offset[0], ymin + label_offset[1])
        roi.label_pos = (xmin + roi.label_offset[0], ymin + roi.label_offset[1])
        roi.label = DraggableRoiLabel(roi, gate_name, self.gates, pos=roi.label_pos)
        self.vb.addItem(roi.label)
        self.rois[roi] = None
        self.update_roi_label_and_gate(roi, create=True)
        roi.sigRegionChanged.connect(lambda: self.update_roi_label_and_gate(roi))

    def new_rectangle_gate(self):
        # set name
        name_increment = 1
        while f'R{name_increment}' in [g[0] for g in self.gates.get_gate_ids()]:
            name_increment += 1
        gate_name = f'R{name_increment}'

        # create ROI
        x0 = 0.35 + 0.05 * name_increment
        y0 = 0.45 - 0.05 * name_increment
        Dx = 0.2
        Dy = 0.2
        roi = BigHandlesRectROI(
            [x0, y0], [Dx, Dy],
            pen=pg.mkPen('g', width=3),
            movable=True,
            removable=True)
        roi.gate_name = gate_name
        roi.sigRemoveRequested.connect(self.remove_gate)
        self.vb.addItem(roi)

        # create label
        roi.label_offset = (label_offset[0], label_offset[1])
        roi.label_pos = (x0 + roi.label_offset[0], y0 + roi.label_offset[1])
        roi.label = DraggableRoiLabel(roi, gate_name, self.gates, pos=roi.label_pos)
        self.vb.addItem(roi.label)
        self.rois[roi] = None
        self.update_roi_label_and_gate(roi, create=True)
        roi.sigRegionChanged.connect(lambda: self.update_roi_label_and_gate(roi))

    def new_ellipse_gate(self):
        # set name
        name_increment = 1
        while f'E{name_increment}' in [g[0] for g in self.gates.get_gate_ids()]:
            name_increment += 1
        gate_name = f'E{name_increment}'

        # create ROI
        x0 = 0.3 + 0.05 * name_increment
        y0 = 0.4 - 0.05 * name_increment
        Dx = 0.3
        Dy = 0.2
        roi = BigHandlesEllipseROI(
            [x0, y0], [Dx, Dy],
            pen=pg.mkPen('g', width=3),
            movable=True,
            removable=True)
        roi.gate_name = gate_name
        roi.sigRemoveRequested.connect(self.remove_gate)
        self.vb.addItem(roi)

        # create label
        roi.label_offset = (label_offset[0], label_offset[1])
        roi.label_pos = (x0 + roi.label_offset[0], y0 + roi.label_offset[1])
        roi.label = DraggableRoiLabel(roi, gate_name, self.gates, pos=roi.label_pos)
        self.vb.addItem(roi.label)
        self.rois[roi] = None
        self.update_roi_label_and_gate(roi, create=True)
        roi.sigRegionChanged.connect(lambda: self.update_roi_label_and_gate(roi))

    def new_range_gate(self):
        # set name
        name_increment = 1
        while f'R_{name_increment}' in [g[0] for g in self.gates.get_gate_ids()]:
            name_increment += 1
        gate_name = f'R_{name_increment}'

        # create ROI
        x1 = 0.35
        x2 = 0.65
        roi = RangeROI(self.vb, x1=x1, x2=x2, removable=True)
        roi.gate_name = gate_name
        roi.sigRemoveRequested.connect(self.remove_gate)
        #roi.sigRangeChanged.connect(lambda x1, x2: print(f"Gate moved to: x1={x1:.2f}, x2={x2:.2f}"))

        # create label
        roi.label_offset = (x1 + label_offset[0], 1 + label_offset[1])
        roi.label_pos = (x1 + roi.label_offset[0], 1 + roi.label_offset[1])
        roi.label = DraggableRoiLabel(roi, gate_name, self.gates, pos=roi.label_pos, anchor=(0,1))
        self.vb.addItem(roi.label)
        self.rois[roi] = None
        self.update_roi_label_and_gate(roi, create=True)
        roi.sigRangeChanged.connect(lambda: self.update_roi_label_and_gate(roi))

    def new_quadrant_gate(self):
        # set name
        name_increment = 1
        while f'Q{name_increment}' in [g[0] for g in self.gates.get_gate_ids()]:
            name_increment += 1
        gate_name = f'Q{name_increment}'

        # create ROI
        x = 0.5
        y = 0.5
        roi = QuadROI(self.vb, x=x, y=y)
        roi.gate_name = gate_name
        roi.sigRemoveRequested.connect(self.remove_gate)
        #quadroi.addToVb(self.vb)
        #quadroi.sigGateMoved.connect(lambda x, y: print(f"Gate moved to: x={x:.2f}, y={y:.2f}"))

        # create labels
        xlim = self.limits_map[self.id_channel_x]
        ylim = self.limits_map[self.id_channel_y]
        roi.label = [
            DraggableRoiLabel(roi, gate_name+'++', self.gates, pos=(xlim[1],ylim[1]), anchor=(1,0)),
            DraggableRoiLabel(roi, gate_name+'+-', self.gates, pos=(xlim[1],ylim[0]), anchor=(1,1)),
            DraggableRoiLabel(roi, gate_name+'-+', self.gates, pos=(xlim[0],ylim[1]), anchor=(0,0)),
            DraggableRoiLabel(roi, gate_name+'--', self.gates, pos=(xlim[0],ylim[0]), anchor=(0,1)),
        ]
        [self.vb.addItem(label) for label in roi.label]
        self.rois[roi] = None
        self.update_roi_label_and_gate(roi, create=True)
        roi.sigPosChanged.connect(lambda: self.update_roi_label_and_gate(roi))

    def update_roi_label_and_gate(self, roi, create=False):
        # get gate name first
        gate_name = roi.gate_name

        # update roi label position - all except quadrant roi
        if roi.label is not None and type(roi.label) is not list and not roi.label.isUnderMouse() :
            roi_pos = roi.pos()
            new_label_pos = (roi_pos.x() + roi.label_offset[0],
                             roi_pos.y() + roi.label_offset[1])
            print(f'new_label_pos{new_label_pos}')
            roi.label.setPos(*new_label_pos)

        # roi.setState({'pos': QtCore.QPointF(0.350000, 0.350000), 'size': QtCore.QPointF(0.300000, 0.200000), 'angle': 45.0})
        if isinstance(roi, QuadROI):
            x = roi.vx.value()
            y = roi.vy.value()

            # QuadrantDivider instances are similar to a Dimension, they take compensation_ref and tranformation_ref
            quad_div_x = fk.QuadrantDivider('xdiv', self.channel_labels[self.id_channel_x], compensation_ref='uncompensated',
                                 transformation_ref=self.channel_labels[self.id_channel_x], values=[x])
            quad_div_y = fk.QuadrantDivider('ydiv', self.channel_labels[self.id_channel_y], compensation_ref='uncompensated',
                                 transformation_ref=self.channel_labels[self.id_channel_y], values=[y])

            quad_divs = [quad_div_x, quad_div_y]

            # the 2 dividers above will be used to divide the space into 4 quadrants
            quad_pp = fk.gates.Quadrant(
                quadrant_id=f'{self.channel_labels[self.id_channel_x]}+ {self.channel_labels[self.id_channel_y]}+',
                divider_refs=['xdiv','ydiv'],
                divider_ranges=[(x, None), (y, None)]
            )
            quad_pn = fk.gates.Quadrant(
                quadrant_id=f'{self.channel_labels[self.id_channel_x]}+ {self.channel_labels[self.id_channel_y]}-',
                divider_refs=['xdiv','ydiv'],
                divider_ranges=[(x, None), (None, y)]
            )
            quad_np = fk.gates.Quadrant(
                quadrant_id=f'{self.channel_labels[self.id_channel_x]}- {self.channel_labels[self.id_channel_y]}+',
                divider_refs=['xdiv','ydiv'],
                divider_ranges=[(None, x), (y, None)]
            )
            quad_nn = fk.gates.Quadrant(
                quadrant_id=f'{self.channel_labels[self.id_channel_x]}- {self.channel_labels[self.id_channel_y]}-',
                divider_refs=['xdiv','ydiv'],
                divider_ranges=[(None, x), (None, y)]
            )
            quadrants = [quad_pp, quad_pn, quad_np, quad_nn]

            if create:
                # define gate
                gate = fk.gates.QuadrantGate(
                    gate_name,
                    dividers=quad_divs,
                    quadrants=quadrants
                )
                [label.labelRenamed.connect(lambda new_name: self.rename_gate(gate, roi, new_name)) for label in roi.label]

                # add gate to hierarchy
                if self.id_source_gate == 0:
                    gate_path = ('root',)
                else:
                    gate_paths = [('root',)] + self.gates.get_gate_ids()
                    gate_id = gate_paths[self.id_source_gate]
                    gate_path = gate_id[1] + (gate_id[0],)
                self.gates.add_gate(gate, gate_path=gate_path)
                self.rois[roi] = (gate_name, gate_path)

            else:
                gate = self.gates.get_gate(gate_name)
                gate.quadrants = {q.id: q for q in quadrants}

            # print([gate, gate.quadrants])
            print([q[1]._divider_ranges for q in gate.quadrants.items()])

        elif isinstance(roi, RangeROI):
            x1 = roi.v1.value()
            x2 = roi.v2.value()

            dim_x = fk.Dimension(self.channel_labels[self.id_channel_x], compensation_ref='uncompensated',
                                 transformation_ref=self.channel_labels[self.id_channel_x], range_min=x1,
                                 range_max=x2)
            if create:
                # define gate
                gate = fk.gates.RectangleGate(gate_name, dimensions=[dim_x])
                roi.label.labelRenamed.connect(lambda new_name: self.rename_gate(gate, roi, new_name))

                # add gate to hierarchy
                if self.id_source_gate == 0:
                    gate_path = ('root',)
                else:
                    gate_paths = [('root',)] + self.gates.get_gate_ids()
                    gate_id = gate_paths[self.id_source_gate]
                    gate_path = gate_id[1] + (gate_id[0],)
                self.gates.add_gate(gate, gate_path=gate_path)
                self.rois[roi] = (gate_name, gate_path)

            else:
                gate = self.gates.get_gate(gate_name)
                gate.dimensions = [dim_x]

            print([gate, gate.dimensions, gate.dimensions[0].min, gate.dimensions[0].max])

        elif isinstance(roi, BigHandlesPolyLineROI):
            roi_state = roi.getState()
            origin = roi_state['pos']
            vertices = roi_state['points']
            #vertices = [h['item'].pos() for h in roi.handles]
            vertices = [(origin.x() + v.x(), origin.x() + v.y()) for v in vertices]

            if create:
                # define gate
                dim_x = fk.Dimension(self.channel_labels[self.id_channel_x], compensation_ref='uncompensated',
                                     transformation_ref=self.channel_labels[self.id_channel_x], range_min=0, range_max=1)
                dim_y = fk.Dimension(self.channel_labels[self.id_channel_y], compensation_ref='uncompensated',
                                     transformation_ref=self.channel_labels[self.id_channel_y], range_min=0, range_max=1)

                gate = fk.gates.PolygonGate(gate_name, [dim_x, dim_y], vertices, use_complement=False)
                roi.label.labelRenamed.connect(lambda new_name: self.rename_gate(gate, roi, new_name))

                # add gate to hierarchy
                if self.id_source_gate == 0:
                    gate_path = ('root',)
                else:
                    gate_paths = [('root',)] + self.gates.get_gate_ids()
                    gate_id = gate_paths[self.id_source_gate]
                    gate_path = gate_id[1] + (gate_id[0],)
                self.gates.add_gate(gate, gate_path=gate_path)
                self.rois[roi] = (gate_name, gate_path)

            else:
                gate = self.gates.get_gate(gate_name)
                gate.vertices = vertices

            print([gate, gate.vertices])

        elif isinstance(roi, BigHandlesRectROI):
            roi_state = roi.getState()
            x0, y0 = np.array(roi_state['pos'])
            Dx, Dy = np.array(roi_state['size']) / 2

            dim_x = fk.Dimension(self.channel_labels[self.id_channel_x], compensation_ref='uncompensated',
                                 transformation_ref=self.channel_labels[self.id_channel_x], range_min=x0,
                                 range_max=x0 + Dx)
            dim_y = fk.Dimension(self.channel_labels[self.id_channel_y], compensation_ref='uncompensated',
                                 transformation_ref=self.channel_labels[self.id_channel_y], range_min=y0,
                                 range_max=y0 + Dy)

            if create:
                # define gate
                gate = fk.gates.RectangleGate(gate_name, dimensions=[dim_x, dim_y])
                roi.label.labelRenamed.connect(lambda new_name: self.rename_gate(gate, roi, new_name))

                # add gate to hierarchy
                if self.id_source_gate == 0:
                    gate_path = ('root',)
                else:
                    gate_paths = [('root',)] + self.gates.get_gate_ids()
                    gate_id = gate_paths[self.id_source_gate]
                    gate_path = gate_id[1] + (gate_id[0],)
                self.gates.add_gate(gate, gate_path=gate_path)
                self.rois[roi] = (gate_name, gate_path)

            else:
                gate = self.gates.get_gate(gate_name)
                gate.dimensions = [dim_x, dim_y]

            #print([gate, gate.dimensions, gate.dimensions[0].min, gate.dimensions[0].max, gate.dimensions[1].min, gate.dimensions[1].max])


        elif isinstance(roi, BigHandlesEllipseROI):
            roi_state = roi.getState()
            coordinates = roi_state['pos'] + roi_state['size'] / 2
            w, h = np.array(roi_state['size']) / 2
            theta = np.deg2rad(roi_state['angle'])

            # Covariance matrix
            R = np.array([[np.cos(theta), -np.sin(theta)],
                          [np.sin(theta), np.cos(theta)]])
            D = np.diag([w ** 2, h ** 2])
            covariance_matrix = R @ D @ R.T
            distance_square = w * h

            if create:
                # define gate
                dim_x = fk.Dimension(self.channel_labels[self.id_channel_x], compensation_ref='uncompensated',
                                     transformation_ref=self.channel_labels[self.id_channel_x], range_min=0, range_max=1)
                dim_y = fk.Dimension(self.channel_labels[self.id_channel_y], compensation_ref='uncompensated',
                                     transformation_ref=self.channel_labels[self.id_channel_y], range_min=0, range_max=1)

                gate = fk.gates.EllipsoidGate(gate_name, [dim_x, dim_y], coordinates, covariance_matrix, distance_square)
                roi.label.labelRenamed.connect(lambda new_name: self.rename_gate(gate, roi, new_name))

                # add gate to hierarchy
                if self.id_source_gate == 0:
                    gate_path = ('root',)
                else:
                    gate_paths = [('root',)] + self.gates.get_gate_ids()
                    gate_id = gate_paths[self.id_source_gate]
                    gate_path = gate_id[1] + (gate_id[0],)
                self.gates.add_gate(gate, gate_path=gate_path)
                self.rois[roi] = (gate_name, gate_path)

            else:
                gate = self.gates.get_gate(gate_name)
                gate.coordinates = coordinates
                gate.covariance_matrix = covariance_matrix
                gate.distance_square = distance_square

            #print([gate, gate.coordinates, gate.covariance_matrix, gate.distance_square])

        self.configure_plot_title_menu()
        print(self.rois)
        print(self.gates.get_gate_ids())

    def rename_gate(self, gate, roi, new_name):
        if self.gates.find_matching_gate_paths(new_name):
            raise Exception(f"gate name {new_name} already exists")
        else:
            self.gates.rename_gate(gate.gate_name, new_name)
            (gate_name, gate_path) = self.rois[roi]
            self.rois[roi] = (new_name, gate_path)
            roi.gate_name = new_name

        print(self.gates.get_gate_ids())
        print(self.rois)

    def remove_gate(self):
        """Remove ROI and associated gate"""
        roi = self.sender()
        if roi is not None:
            (gate_name, gate_path) = self.rois[roi]
            self.gates.remove_gate(gate_name, gate_path=gate_path, keep_children=True)

            self.vb.removeItem(roi.label)
            self.vb.removeItem(roi)
            self.rois.pop(roi) # or del self.rois[roi]

            # update menus
            self.configure_axes_labels_menus()

            print(self.rois)
            print(self.gates.get_gate_ids())

    def apply_zoom(self, axisname):
        if axisname == 'x':
            axis = self.axis_bottom
        elif axisname == 'y':
            axis = self.axis_left

        if axis._pending_delta == 0:
            return

        # Accumulate small changes
        threshold = 1  # pixels
        step = axis._pending_delta
        axis._pending_delta = 0
        zoom_rate = 1.04  # tune this

        if abs(step) < threshold:
            return

        if step > 0:
            factor = 1 / zoom_rate
        else:
            factor = zoom_rate



        if axisname == 'x':

            x_min, x_max = axis.vb.viewRange()[0]
            if self.transform_map[self.id_channel_x] == 0 or self.transform_map[self.id_channel_x] == 2: # linear or log
                new_xmax = (x_max - axis.zoomZero) * factor + axis.zoomZero
                new_xmin = (x_min - axis.zoomZero) * factor + axis.zoomZero
                if new_xmax < axis.fullRange[1] * 1.01:
                    axis.vb.setXRange(new_xmin, new_xmax, padding=0)
                axis.limits = (new_xmin, new_xmax)
                self.limits_map[self.id_channel_x] = axis.limits
                self.transform_x.set_transform(self.transform_map[self.id_channel_x],
                                               self.limits_map[self.id_channel_x])

            elif self.transform_map[self.id_channel_x] == 1:
                # scale w in bottom half, limits in top half
                if self.vb.mapToView(axis.initial_pos).x() < 0.5 * x_max:
                    self.transform_x.logicle_w = self.transform_x.logicle_w / factor
                    self.transform_x.initialise_xforms()
                    self.configure_axes_labels_menus()
                else:
                    new_xmax = (x_max - axis.zoomZero) * factor + axis.zoomZero
                    new_xmin = (x_min - axis.zoomZero) * factor + axis.zoomZero
                    if new_xmax < axis.fullRange[1] * 1.01:
                        axis.vb.setXRange(new_xmin, new_xmax, padding=0)
                    axis.limits = (new_xmin, new_xmax)
                    self.limits_map[self.id_channel_x] = axis.limits
                    self.transform_x.set_transform(self.transform_map[self.id_channel_x],
                                                   self.limits_map[self.id_channel_x])


        elif axisname == 'y':

            y_min, y_max = axis.vb.viewRange()[1]
            if self.transform_map[self.id_channel_y] == 0 or self.transform_map[self.id_channel_y] == 2:
                new_ymax = (y_max - axis.zoomZero) / factor + axis.zoomZero
                new_ymin = (y_min - axis.zoomZero) / factor + axis.zoomZero
                if new_ymax < axis.fullRange[1] * 1.01:
                    axis.vb.setYRange(new_ymin, new_ymax, padding=0)
                axis.limits = (new_ymin, new_ymax)
                self.limits_map[self.id_channel_y] = axis.limits
                self.transform_y.set_transform(self.transform_map[self.id_channel_y],
                                               self.limits_map[self.id_channel_y])

            elif self.transform_map[self.id_channel_y] == 1: # logicle
                # scale w in bottom half, limits in top half
                if self.vb.mapToView(axis.initial_pos).y() < 0.5 * y_max:
                    self.transform_y.logicle_w = self.transform_y.logicle_w * factor
                    self.transform_y.initialise_xforms()
                    self.configure_axes_labels_menus()
                else:
                    new_ymax = (y_max - axis.zoomZero) / factor + axis.zoomZero
                    new_ymin = (y_min - axis.zoomZero) / factor + axis.zoomZero
                    if new_ymax < axis.fullRange[1] * 1.01:
                        axis.vb.setYRange(new_ymin, new_ymax, padding=0)
                    axis.limits = (new_ymin, new_ymax)
                    self.limits_map[self.id_channel_y] = axis.limits
                    self.transform_y.set_transform(self.transform_map[self.id_channel_y],
                                                   self.limits_map[self.id_channel_y])



        if self.event_data_all_channels is not None:
            self.refresh_data()

    def apply_limits(self):
        self.limits_map[self.id_channel_x] = self.axis_bottom.limits
        self.limits_map[self.id_channel_y] = self.axis_left.limits

        # Apply transforms to axes
        self.transform_x.set_transform(self.transform_map[self.id_channel_x], self.limits_map[self.id_channel_x])
        self.transform_y.set_transform(self.transform_map[self.id_channel_y], self.limits_map[self.id_channel_y])

        if self.event_data_all_channels is not None:
            self.calc_and_plot_hist2d()

    def configure_plot_title_menu(self):
        # set title and items for source gate menu
        gate_names = ['root'] + [g[0] for g in self.gates.get_gate_ids()]
        self.plot_title.setText(gate_names[self.id_source_gate])
        self.plot_title.leftClickMenuItems = gate_names
        self.plot_title.leftClickMenuFunction = self.set_source_gate
        self.plot_title.leftItemSelected = self.id_source_gate

    def configure_axes_labels_menus(self):
        self.configure_plot_title_menu()

        # Apply labels to axes, activate heatmap image or 1d histogram
        if self.id_channel_x == 'All Fluorescence':
            self.id_channel_y = 'Intensity'
            self.label_x.setText(self.id_channel_x)
            self.label_y.setText(self.id_channel_y)
            self.img.setVisible(True)
            self.hist.setVisible(False)
        elif self.id_channel_y == 'Count':
            self.label_x.setText(self.channel_labels[self.id_channel_x])
            self.label_y.setText(self.id_channel_y)
            self.img.setVisible(False)
            self.hist.setVisible(True)
            self.vb.enableAutoRange(axis=self.vb.YAxis, enable=True)
        else:
            self.label_x.setText(self.channel_labels[self.id_channel_x])
            self.label_y.setText(self.channel_labels[self.id_channel_y])
            self.img.setVisible(True)
            self.hist.setVisible(False)

        # set items for channel menus
        self.label_x.leftClickMenuItems = self.channel_labels + ['All Fluorescence']
        self.label_x.rightClickMenuItems = transforms_menu_items
        self.label_x.leftClickMenuFunction = self.set_channel_x
        self.label_x.rightClickMenuFunction = self.set_transform_x
        self.label_x.leftItemSelected = self.id_channel_x
        self.label_x.rightItemSelected = self.transform_map[self.id_channel_x]
        self.label_y.leftClickMenuItems = self.channel_labels + ['Count']
        self.label_y.rightClickMenuItems = transforms_menu_items
        self.label_y.leftClickMenuFunction = self.set_channel_y
        self.label_y.rightClickMenuFunction = self.set_transform_y
        self.label_y.leftItemSelected = self.id_channel_y
        self.label_y.rightItemSelected = self.transform_map[self.id_channel_y]

        # Apply transforms to axes
        self.transform_x.set_transform(self.transform_map[self.id_channel_x], self.limits_map[self.id_channel_x])
        self.transform_y.set_transform(self.transform_map[self.id_channel_y], self.limits_map[self.id_channel_y])

        self.transform_x.ticks(self.axis_bottom)
        self.transform_y.ticks(self.axis_left)

        self.axis_bottom.zoomZero = self.transform_x.zero
        self.axis_left.zoomZero = self.transform_y.zero

        # Record transformations in gating strategy
        if type(self.id_channel_x) == int:
            self.gates.transformations[self.channel_labels[self.id_channel_x]] = self.transform_x.xform
        if type(self.id_channel_y) == int:
            self.gates.transformations[self.channel_labels[self.id_channel_y]] = self.transform_y.xform

        # Set limits
        self.axis_bottom.limits = self.limits_map[self.id_channel_x]
        self.vb.setXRange(self.limits_map[self.id_channel_x][0], self.limits_map[self.id_channel_x][1], padding=0)

        if self.id_channel_x == 'All Fluorescence':
            self.vb.enableAutoRange(axis=self.vb.YAxis, enable=True)
            self.axis_bottom.limits = self.limits_map[self.id_channel_x]
            # TODO set x ticks to channel names
            # TODO set y limits and transform to logicle
        elif self.id_channel_y == 'Count':
            self.vb.enableAutoRange(axis=self.vb.YAxis, enable=True)
        else:
            self.axis_left.limits = self.limits_map[self.id_channel_y]
            self.vb.setYRange(self.limits_map[self.id_channel_y][0], self.limits_map[self.id_channel_y][1], padding=0)

        # Refresh data
        if self.event_data_all_channels is not None:
            self.refresh_data()

        export_widget_as_image(self, "widget_screenshot.png")
        export_high_quality(self, "widget_screenshot_hq.png", scale_factor=2)
        render_widget_light_pg(self, "widget_screenshot_light.png", scale_factor=2)
        export_docx("widget_screenshot.png")


    def set_channel_x(self,n):
        self.id_channel_x = n
        self.configure_axes_labels_menus()

    def set_channel_y(self,n):
        self.id_channel_y = n
        self.configure_axes_labels_menus()

    def set_transform_x(self,n):
        self.transform_map[self.id_channel_x] = n
        self.configure_axes_labels_menus()

    def set_transform_y(self,n):
        self.transform_map[self.id_channel_y] = n
        self.configure_axes_labels_menus()

    def set_source_gate(self,n):
        self.id_source_gate = n
        self.configure_axes_labels_menus()

    # def configure_histRibbonPlot(self):
    #     # Choose a Colorcet colormap (e.g., 'fire', 'bgy', 'rainbow')
    #     colormap_name = 'rainbow4'
    #     colors = cc.palette[colormap_name]  # Get the colormap from Colorcet
    #
    #     # Convert Colorcet colormap to PyQtGraph's format
    #     cmap = pg.ColorMap(
    #         pos=np.linspace(0.0, 1.0, len(colors)),
    #         color=colors
    #     )
    #
    #     rgba_lut = cmap.getLookupTable(alpha=True)
    #     rgba_lut[0, 3] = 0  # Fully transparent for 0
    #     self.img.setLookupTable(rgba_lut)
    #     self.vb.addItem(self.img)

    def configure_hist2d(self):
        #cmapname = 'CET-C1'
        #cmapname = 'turbo'  # or 'plasma', 'inferno', etc.
        #cmap = pg.colormap.get(cmapname)

        # Choose a Colorcet colormap (e.g., 'fire', 'bgy', 'rainbow')
        colormap_name = 'rainbow4'
        colors = cc.palette[colormap_name]  # Get the colormap from Colorcet

        # Convert Colorcet colormap to PyQtGraph's format
        cmap = pg.ColorMap(
            pos=np.linspace(0.0, 1.0, len(colors)),
            color=colors
        )

        rgba_lut = cmap.getLookupTable(alpha=True)
        rgba_lut[0, 3] = 0  # Fully transparent for 0
        self.img.setLookupTable(rgba_lut)
        self.vb.addItem(self.img)

    def configure_hist1d(self):
        # Add curve to the existing ViewBox
        self.vb.addItem(self.hist)
        self.count = np.zeros(self.transform_x.scale_bins+2)

    def refresh_data(self):
        if self.id_channel_x == 'All Fluorescence':
            self.calc_and_plot_ribbon_plot()
        elif self.id_channel_y == 'Count':
            self.calc_and_plot_hist1d()
        else:
            self.calc_and_plot_hist2d()

    def calc_and_plot_ribbon_plot(self):
        heatmap = np.apply_along_axis(lambda x: np.histogram(x, bins=self.transform_y.scale)[0], axis=0, arr=self.event_data_all_channels[:,self.fluoro_indices])
        self.img.setImage(heatmap.T)

    def calc_and_plot_hist2d(self):

        x = self.event_data_all_channels[:, self.id_channel_x]
        y = self.event_data_all_channels[:, self.id_channel_y]

        start = time.perf_counter()

        # Calculate 2D histogram (density)
        heatmap, xedges, yedges = np.histogram2d(x, y, bins=[self.transform_x.scale, self.transform_y.scale])
        # heatmap = (heatmap - heatmap.min()) / (heatmap.max() - heatmap.min())
        self.img.setImage(heatmap)

        # Set the position and scale of the image
        self.img.setRect(pg.QtCore.QRectF(
            self.limits_map[self.id_channel_x][0],
            self.limits_map[self.id_channel_y][0],
            self.limits_map[self.id_channel_x][1] - self.limits_map[self.id_channel_x][0],
            self.limits_map[self.id_channel_y][1] - self.limits_map[self.id_channel_y][0]
        ))

        end = time.perf_counter()
        print(f"Execution time: {end - start:.6f} seconds")

    def calc_and_plot_hist1d(self):
        x = self.event_data_all_channels[:, self.id_channel_x]

        # Calculate 1D histogram
        self.count[:-1], xedges = np.histogram(x, bins=self.transform_x.scale)
        self.hist.setData(self.transform_x.step_scale, self.count)


class TestWindow(QtWidgets.QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("PySide6 with Interactive Cytometry Plot")
        self.setGeometry(100, 100, 600, 600)

        # Create central widget
        central_widget = QtWidgets.QWidget()
        self.setCentralWidget(central_widget)

        # Create layout
        layout = QtWidgets.QVBoxLayout(central_widget)

        # win = pg.GraphicsLayoutWidget()
        # win.resize(600, 600)

        fcs_path = '../../flowkit_demos/FlowKit-master/data/8_color_data_set/fcs_files/101_DEN084Y5_15_E01_008_clean.fcs'
        sample = fk.Sample(fcs_path)
        np_events = sample.get_events(source='raw')
        channel_labels = sample.pnn_labels
        scatter_indices = sample.scatter_indices
        fluoro_indices = sample.fluoro_indices
        gating_strategy = fk.GatingStrategy()

        transform_map = {}
        limits_map = {}

        print(sample.get_metadata())
        print(sample.channels)
        print(sample.fluoro_indices)
        print(sample.scatter_indices)
        print(sample.time_index)
        print(sample.acquisition_date)
        print(sample.compensation)
        print(sample.transform)
        print(sample.event_count)

        self.plot_widget = InteractiveCytometryPlotWidget()
        layout.addWidget(self.plot_widget)

        self.plot_widget.configure_plot(channel_labels, scatter_indices, fluoro_indices, transform_map, limits_map, gating_strategy)
        self.plot_widget.event_data_all_channels = np_events
        self.plot_widget.refresh_data()


if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)

    win = TestWindow()
    win.show()
    sys.exit(app.exec())
