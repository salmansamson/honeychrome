import io
import time
from datetime import datetime
from pathlib import Path

import numpy as np
from PySide6.QtCore import QObject, QEventLoop, Signal
from PySide6.QtGui import QPalette, QColor
from PySide6.QtWidgets import QApplication, QWidget, QHeaderView
import pyqtgraph as pg
import docx
from docx.shared import Mm, Inches, Pt
from docx.oxml import parse_xml

from honeychrome.settings import cytometry_plot_width_export
import honeychrome.settings as settings
from honeychrome.view_components.busy_cursor import with_busy_cursor
from honeychrome.view_components.cytometry_plot_widget import CytometryPlotWidget, pm_to_png_buffer, get_widget_pixmap, export_widget_png
from honeychrome.view_components.heatmap_viewedit import HeatmapViewEditor
from honeychrome.view_components.nxn_grid import NxNGrid
from honeychrome.view_components.profiles_viewer import ProfilesViewer


def resize_tableview(widget):
    table = widget.view

    # Get sizes
    width  = table.verticalHeader().width() + sum(table.columnWidth(i) for i in range(table.model().columnCount()))
    height = table.horizontalHeader().height() + sum(table.rowHeight(i) for i in range(table.model().rowCount()))

    # Add frame width (for borders)
    frame = table.frameWidth() * 2

    # Final dimensions
    total_width  = width + frame
    total_height = height + frame

    # Set parent size
    parent = table.parentWidget()
    parent.setFixedSize(total_width, total_height)


def set_widget_light_palette(widget):
    palette = widget.palette()

    # Set light theme colors
    palette.setColor(QPalette.Window, QColor(255, 255, 255))
    palette.setColor(QPalette.WindowText, QColor(0, 0, 0))
    palette.setColor(QPalette.Base, QColor(255, 255, 255))
    palette.setColor(QPalette.Text, QColor(0, 0, 0))
    palette.setColor(QPalette.Button, QColor(255, 255, 255))
    palette.setColor(QPalette.ButtonText, QColor(0, 0, 0))

    widget.setPalette(palette)
    widget.setAutoFillBackground(True)  # Important!


def add_inline_picture(doc, png_buffer, own_line):
    """
    Add an inline picture that behaves like a character without adding newline
    """
    # Create a paragraph if needed
    if not hasattr(doc, '_current_run') or doc._current_run is None:
        doc._current_paragraph = doc.add_paragraph()
        doc._current_run = doc._current_paragraph.add_run()

    # Get the run
    run = doc._current_run

    # Add the picture
    if own_line:
        doc._current_paragraph = doc.add_paragraph()
        doc._current_run = doc._current_paragraph.add_run()
        picture = run.add_picture(png_buffer, width=Mm(170))
        doc._current_paragraph = doc.add_paragraph()
        doc._current_run = doc._current_paragraph.add_run()
    else:
        picture = run.add_picture(png_buffer, width=Mm(cytometry_plot_width_export))

    # Remove the paragraph properties that cause newlines
    # by ensuring the picture is truly inline
    return picture

def reliable_table_autofit(table):
    """Always use this for reliable auto-fit"""

    tbl_pr = table._element.tblPr

    # Clear existing layout
    for elem in tbl_pr.xpath('.//w:tblLayout'):
        tbl_pr.remove(elem)

    # Add autofit layout
    autofit_xml = r'<w:tblLayout w:type="autofit" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    tbl_pr.append(parse_xml(autofit_xml))

    # Also set the python-docx properties
    table.autofit = True
    table.allow_autofit = True

def format_table(table):
    # Make entire row 1 (second row, index 1) bold
    row_index = 0  # Row to make bold (0-based index)
    row = table.rows[row_index]

    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

class ReportGenerator(QObject):
    finished = Signal()

    def __init__(self, bus, controller):
        super().__init__()
        self.bus = bus
        self.controller = controller

    @with_busy_cursor
    def export(self):
        current_sample = self.controller.current_sample
        print(f'ReportGenerator: started {current_sample}')

        # Do some checks
        if not current_sample:
            self.bus.warningMessage.emit('Please select a sample before generating a report')
            return
        if self.controller.current_sample.event_count == 0:
            self.bus.warningMessage.emit('Sample has no events')
            return

        self.bus.statusMessage.emit(f'ReportGenerator: started {current_sample}')

        current_sample_path = Path(self.controller.current_sample_path)
        report_rel_path = Path('Reports') / current_sample_path.relative_to(self.controller.experiment.settings['raw']['raw_samples_subdirectory'])
        full_sample_path = self.controller.experiment_dir / report_rel_path
        full_sample_path.parent.mkdir(parents=True, exist_ok=True)
        docx_path = str(full_sample_path.with_suffix('.docx'))
        current_mode = self.controller.current_mode
        resolution = 450
        scale_factor = cytometry_plot_width_export / 25.4 * 300 / resolution # for 300 DPI, width in mm

        # --- Save old PyQtGraph colors ---
        old_bg = pg.getConfigOption("background")
        old_fg = pg.getConfigOption("foreground")

        # set up document
        doc = docx.Document()
        section = doc.sections[0]
        margin = Mm(15)
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin
        style = doc.styles['Normal']
        font = style.font
        # font.name = 'sans'
        font.size = Pt(8)

        doc.add_heading('Honeychrome Report', 1)
        doc.add_heading(str(current_sample_path.stem), 2)
        desc = (f"Experiment: {str(self.controller.experiment_dir)}\n"
                f"Sample: {str(current_sample_path)}\n"
                f"Event count: {current_sample.event_count}\n"
                f'Report generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S %Z')}\n')
        try:
            desc += f"Sample acquired: {current_sample.get_metadata()['date']}; begin time: {current_sample.get_metadata()['btim']}; end time: {current_sample.get_metadata()['etim']}"
        except:
            print('No acquisition date in metadata')
        doc.add_paragraph(desc)

        self.bus.statusMessage.emit(f'ReportGenerator: added header')

        # Cytometry plots and stats
        mode_tab = []
        container = QWidget()
        pg.setConfigOptions(background='w', foreground='k')

        if settings.report_include_raw_retrieved:
            mode_tab.append(('raw', 'Raw Data'))
        if self.controller.experiment.process['unmixing_matrix'] and settings.report_include_unmixed_retrieved:
            mode_tab.append(('unmixed', 'Unmixed Data'))

        for mode, tab in mode_tab:
            self.bus.statusMessage.emit(f'ReportGenerator: adding {tab}')

            self.controller.set_mode(tab)
            time.sleep(1)
            doc.add_heading(tab, 3)

            doc._current_paragraph = doc.add_paragraph()
            doc._current_run = doc._current_paragraph.add_run()
            for n, plot in enumerate(self.controller.data_for_cytometry_plots['plots']):
                self.bus.statusMessage.emit(f'ReportGenerator: adding plot {n}')
                plot_widget = CytometryPlotWidget(mode=mode, n_in_plot_sequence=n, plot=plot, data_for_cytometry_plots=self.controller.data_for_cytometry_plots, parent=container)
                if plot['type'] == 'ribbon':
                    height = 1.5 * resolution
                    width = 1.5 * 2.5 * resolution
                    own_line = True
                else:
                    height = resolution
                    width = resolution
                    own_line = False

                plot_widget.setFixedSize(width, height)
                plot_widget.plot_histogram()
                for roi in plot_widget.rois:
                    roi.label.add_statistic_to_name()
                png_buffer = pm_to_png_buffer(get_widget_pixmap(plot_widget, scale_factor=scale_factor))
                add_inline_picture(doc, png_buffer, own_line)

                ### export also all individual plots
                # export_widget_png(plot_widget, full_sample_path.with_name(f"{current_sample_path.stem}_{plot_widget.n_in_plot_sequence}").with_suffix('.docx'), scale_factor=scale_factor)

            self.bus.statusMessage.emit(f'ReportGenerator: adding statistics')
            statistics = self.controller.data_for_cytometry_plots['statistics']
            table = doc.add_table(rows=1, cols=5)
            row = table.rows[0]
            row.cells[0].text = 'Gate'
            row.cells[1].text = 'N Events'
            row.cells[2].text = '% Total'
            row.cells[3].text = '% Parent'
            row.cells[4].text = 'Conc [/uL]'
            for gate in statistics.keys():
                stats_line = statistics[gate]
                row = table.add_row()
                row.cells[0].text = gate
                row.cells[1].text = f'{stats_line['n_events_gate']}'
                row.cells[2].text = f'{stats_line['p_gate_total']*100:0.2f}'
                row.cells[3].text = f'{stats_line['p_gate_parent']*100:0.2f}'
                row.cells[4].text = f'{stats_line['event_conc']*100:0.2f}' if not np.isnan(stats_line['event_conc']) else ''
            reliable_table_autofit(table)
            format_table(table)

        if self.controller.experiment.process['unmixing_matrix'] and settings.report_include_process_retrieved:
            self.bus.statusMessage.emit(f'ReportGenerator: adding spectral process')
            self.controller.set_mode('Spectral Process')
            doc.add_heading('Spectral Model', 3)
            doc.add_paragraph(f"Negative Type: {self.controller.experiment.process['negative_type']}")
            table = doc.add_table(rows=1, cols=5)
            row = table.rows[0]
            row.cells[0].text = 'Label'
            row.cells[1].text = 'Control Type'
            row.cells[2].text = 'Particle Type'
            row.cells[3].text = 'Gate Channel'
            row.cells[4].text = 'Gate Label'
            for control in self.controller.experiment.process['spectral_model']:
                row = table.add_row()
                row.cells[0].text = control['label'] if control['label'] else ''
                row.cells[1].text = control['control_type'] if control['control_type'] else ''
                row.cells[2].text = control['particle_type'] if control['particle_type'] else ''
                row.cells[3].text = control['gate_channel'] if control['gate_channel'] else ''
                row.cells[4].text = control['gate_label'] if control['gate_label'] else ''
            reliable_table_autofit(table)
            format_table(table)

            self.bus.statusMessage.emit(f'ReportGenerator: adding profiles')
            doc.add_heading('Profiles', 3)
            profiles_viewer = ProfilesViewer(None, self.controller, pen_width=5)
            profiles_viewer.title.setVisible(False)
            set_widget_light_palette(profiles_viewer)
            profiles_viewer.setFixedSize(resolution * 4, resolution * 2)
            png_buffer = pm_to_png_buffer(get_widget_pixmap(profiles_viewer, scale_factor=scale_factor))
            doc.add_picture(png_buffer, width=Mm(170))

            self.bus.statusMessage.emit(f'ReportGenerator: adding similarity matrix')
            doc.add_heading('Similarity Matrix', 3)
            similarity_viewer = HeatmapViewEditor(None, self.controller, 'similarity_matrix', False, parent=container)
            similarity_viewer.title.setVisible(False)
            set_widget_light_palette(similarity_viewer)
            resize_tableview(similarity_viewer)
            png_buffer = pm_to_png_buffer(get_widget_pixmap(similarity_viewer, scale_factor=scale_factor))
            doc.add_picture(png_buffer, width=Mm(min([170, 14*len(self.controller.experiment.process['spectral_model'])])))

            self.bus.statusMessage.emit(f'ReportGenerator: adding unmixing matrix')
            doc.add_heading('Unmixing Matrix', 3)
            unmixing_viewer = HeatmapViewEditor(None, self.controller, 'unmixing_matrix', False, parent=container)
            unmixing_viewer.title.setVisible(False)
            set_widget_light_palette(unmixing_viewer)
            resize_tableview(unmixing_viewer)
            png_buffer = pm_to_png_buffer(get_widget_pixmap(unmixing_viewer, scale_factor=scale_factor))
            doc.add_picture(png_buffer, width=Mm(min([170, 14*len(self.controller.filtered_raw_fluorescence_channel_ids)])))

            self.bus.statusMessage.emit(f'ReportGenerator: adding spillover / fine-tuning matrix')
            doc.add_heading('Spillover / Fine-Tuning Matrix', 3)
            compensation_editor = HeatmapViewEditor(None, self.controller, 'spillover', False, parent=container)
            compensation_editor.title.setVisible(False)
            set_widget_light_palette(compensation_editor)
            resize_tableview(compensation_editor)
            png_buffer = pm_to_png_buffer(get_widget_pixmap(compensation_editor, scale_factor=scale_factor))
            doc.add_picture(png_buffer, width=Mm(min([170, 14*len(self.controller.experiment.process['spectral_model'])])))

            self.bus.statusMessage.emit(f'ReportGenerator: adding NxN plots')
            doc.add_heading('NxN Plots', 3)
            nxn_viewer = NxNGrid(None, self.controller, is_dark=False, parent=container)
            nxn_viewer.title.setVisible(False)
            nxn_viewer.help_nxn.setVisible(False)
            nxn_viewer.source_gate_combo.setVisible(False)
            set_widget_light_palette(nxn_viewer)
            resize_tableview(nxn_viewer)
            png_buffer = pm_to_png_buffer(get_widget_pixmap(nxn_viewer, scale_factor=scale_factor))
            doc.add_picture(png_buffer, width=Mm(min([170, 30*len(self.controller.experiment.process['spectral_model'])])))

            profiles_viewer.deleteLater()
            similarity_viewer.deleteLater()
            unmixing_viewer.deleteLater()
            compensation_editor.deleteLater()
            nxn_viewer.deleteLater()

        # Save - one line!
        doc.save(docx_path)
        container.deleteLater()

        # restore cytometry
        self.controller.set_mode(current_mode)
        pg.setConfigOptions(background=old_bg, foreground=old_fg)

        print(f'ReportGenerator: finished {docx_path}')
        self.bus.statusMessage.emit(f'ReportGenerator: exported {docx_path}')
        self.bus.popupMessage.emit(f"Exported report: {docx_path}")
        self.finished.emit()

